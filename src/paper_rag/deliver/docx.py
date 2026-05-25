"""Word document generator (M10 / ADR-0016).

Converts a multi-paper survey into a formatted .docx using ``python-docx``.
The structure mirrors the Markdown survey (Introduction / Methods Comparison
/ Open Problems / References) but with proper Word styles (Heading 1/2,
Quote block, numbered list for refs).

Two-step pipeline
-----------------
1. Reuse ``survey_md.generate()`` to get the canonical Markdown body.
2. Parse the Markdown headings / paragraphs / refs and emit corresponding
   Word elements. We do NOT use a Markdown-to-Word library — Markdown here
   is a constrained subset (only `## H2`, paragraphs, `### H3`, `- bullets`)
   so a 50-line parser handles it cleanly and avoids the
   `markdown-it-py + html2docx` dep stack.
"""

from __future__ import annotations

import datetime as dt
import io
import logging
import re

from .dispatch import DeliverableResult
from . import survey_md

log = logging.getLogger(__name__)


def _new_doc():
    """Lazy import; raise informative error if python-docx missing."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for docx delivery. "
            "Install with: pip install -e .[deliver]"
        ) from e
    return Document()


_INLINE_CITE_RE = re.compile(r"\[chunk:([0-9a-f]{8,})\]")


def _strip_cites(text: str) -> tuple[str, list[str]]:
    """Pull `[chunk:xxx]` out of body text into a footnote-friendly list.

    For Word output we keep cites inline (looks fine) but also collect them
    so callers can rebuild a refs list. Returns (text, [chunk_ids]).
    """
    cids = _INLINE_CITE_RE.findall(text)
    return text, cids


def _md_to_docx(doc, md: str) -> None:
    """Translate our constrained Markdown subset into Word elements.

    Supported Markdown features:
      # H1     → Heading 1
      ## H2    → Heading 2
      ### H3   → Heading 3
      > quote  → Intense Quote style
      - bullet → bullet list
      blank line → paragraph break
      otherwise → regular paragraph
    """
    bullet_state = False
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            bullet_state = False
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            bullet_state = False
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            bullet_state = False
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            bullet_state = False
            continue
        if line.startswith("> "):
            try:
                doc.add_paragraph(line[2:].strip(), style="Intense Quote")
            except KeyError:
                doc.add_paragraph(line[2:].strip())
            bullet_state = False
            continue
        if line.startswith("- "):
            try:
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            except KeyError:
                doc.add_paragraph("• " + line[2:].strip())
            bullet_state = True
            continue
        # numbered list (1. xxx)
        m = re.match(r"^\d+\.\s+(.*)$", line)
        if m:
            try:
                doc.add_paragraph(m.group(1), style="List Number")
            except KeyError:
                doc.add_paragraph(m.group(0))
            bullet_state = True
            continue
        # default: regular paragraph
        text, _cids = _strip_cites(line.strip())
        doc.add_paragraph(text)
        bullet_state = False


def generate(
    paper_ids: list[str],
    *,
    title: str | None = None,
    max_words: int = 5000,
    **kwargs,
) -> DeliverableResult:
    """Generate a .docx deliverable.

    Internally regenerates the Markdown survey then translates it. Cost is
    one extra Markdown serialization pass over the same data — no second
    LLM call.
    """
    title = title or "Literature Survey"

    md_result = survey_md.generate(
        paper_ids,
        title=title,
        max_words=max_words,
    )
    md_text = md_result.content_bytes.decode("utf-8")

    doc = _new_doc()
    _md_to_docx(doc, md_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    today = dt.date.today().isoformat()
    safe_title = title.lower().replace(" ", "_").replace("/", "_")[:60]

    metadata = dict(md_result.metadata)
    metadata["source_format"] = "markdown_survey"
    metadata["source_size_bytes"] = len(md_result.content_bytes)

    return DeliverableResult(
        format="docx",
        filename=f"{safe_title}_{today}.docx",
        content_bytes=buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        metadata=metadata,
    )


__all__ = ["generate"]
